"""
ingestion_pipeline.py
---------------------
Phase 2: Document Ingestion

Loads all staged .docx documents from data/staging/, extracts text using a
table-aware extractor, cleans Unicode artifacts, and passes serialized text
to the chunker (Phase 3).

Key Features:
  - Table-aware extraction: serializes sub-process table blocks into labeled text.
  - Incremental ingestion: skips documents whose SHA-256 hash is unchanged.
  - Duplicate detection: rejects uploads identical to already-indexed documents.
  - Version supersession: marks old document as 'superseded' when a newer version
    of the same module is detected.
  - SQLite auditing: all actions are logged to ingestion_ledger.db.
  - ChromaDB-ready output: produces ChunkRecord.to_embedding_payload() dicts.
"""

from __future__ import annotations

import hashlib
import logging
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import docx

import ledger
from chunk_schema import (
    AccessLevel,
    ChunkRecord,
    DocumentCategory,
    DocumentRecord,
    DocumentStatus,
    IngestionSummary,
)
from chunker import run_chunking, get_embedding_payloads

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
STAGING_DIR  = os.path.join(PROJECT_ROOT, "data", "staging")
DATA_DIR     = os.path.join(PROJECT_ROOT, "data")


# ===========================================================================
# SECTION 1: File Hashing
# ===========================================================================

def compute_sha256(filepath: str) -> str:
    """Computes SHA-256 hash of a file for deduplication and change detection."""
    sha256 = hashlib.sha256()
    with open(filepath, "rb") as f:
        while chunk := f.read(8192):
            sha256.update(chunk)
    return sha256.hexdigest()


# ===========================================================================
# SECTION 2: Metadata Extraction Helpers
# ===========================================================================

def _extract_title(paragraphs: list[str]) -> str:
    """Extracts document title from first meaningful paragraph lines."""
    for p in paragraphs:
        p = p.strip()
        if p and len(p) > 5:
            return p[:120]
    return "Untitled"


def _extract_module_name(filename: str, paragraphs: list[str]) -> str:
    """
    Extracts the administrative module/department name.
    Prefers the 'Module: X' header from the document body over filename parsing.
    """
    # Strategy 1: Look for 'Module: X' in first 20 paragraphs
    for p in paragraphs[:20]:
        m = re.match(r'^Module:\s*(.+)$', p.strip(), re.IGNORECASE)
        if m:
            return m.group(1).strip()

    # Strategy 2: Parse from filename — only reliable for the VIT SOP naming
    # convention (a numeric prefix and/or "VIT" in the name, e.g.
    # "10.VIT Research & Development 1.0.docx" → "Research & Development").
    # Documents that don't follow that convention (research papers, arbitrary
    # uploads) would otherwise have their whole filename parsed into a
    # nonsense "department" — fall back to "General" for those instead.
    looks_like_sop_corpus = bool(re.match(r'^\d+[\.\s]', filename)) or bool(
        re.search(r'\bvit\b', filename, re.IGNORECASE)
    )
    if not looks_like_sop_corpus:
        return "General"

    name = os.path.splitext(filename)[0]
    name = re.sub(r'^\d+[\.\s]+', '', name)          # Remove prefix number
    name = re.sub(r'\bvit\b', '', name, flags=re.IGNORECASE)
    name = re.sub(r'\b\d+\.\d+\b', '', name)          # Remove version like 1.0
    name = re.sub(r'\bv\d+(\.\d+)?\b', '', name, flags=re.IGNORECASE)
    name = re.sub(r'\(.*?\)', '', name)                # Remove parenthetical like (3)
    name = re.sub(r'\bfinal\b', '', name, flags=re.IGNORECASE)
    name = re.sub(r'dm$|_hr$', '', name, flags=re.IGNORECASE)
    name = re.sub(r'\.docx?$', '', name, flags=re.IGNORECASE)
    name = re.sub(r'[\s_\-\.]+', ' ', name).strip()
    return name.title() or "General"


def _extract_version(filename: str, paragraphs: list[str]) -> str:
    """Extracts version from filename or body text."""
    # Filename: e.g. '1.0', '(3)', 'Final'
    m = re.search(r'\b(\d+\.\d+)\b', filename)
    if m:
        return m.group(1)
    m = re.search(r'\((\d+)\)', filename)
    if m:
        return f"{m.group(1)}.0"
    if re.search(r'\bfinal\b', filename, re.IGNORECASE):
        return "Final"

    # Body text: 'Version 2.1', 'Rev. 1.3'
    for p in paragraphs[:30]:
        m = re.search(r'(?:version|ver\.|rev\.)\s*(\d+\.\d+)', p, re.IGNORECASE)
        if m:
            return m.group(1)
    return "1.0"


def _extract_date(filepath: str, paragraphs: list[str]) -> str:
    """Extracts issue date from body text, falling back to file modification time."""
    date_patterns = [
        r'\b(\d{4}-\d{2}-\d{2})\b',
        r'\b(\d{2}[/-]\d{2}[/-]\d{4})\b',
        r'\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})\b',
    ]
    body = " ".join(paragraphs[:40])
    for pat in date_patterns:
        m = re.search(pat, body, re.IGNORECASE)
        if m:
            return m.group(1)

    # Filesystem fallback
    try:
        mtime = os.path.getmtime(filepath)
        return datetime.fromtimestamp(mtime).strftime("%Y-%m-%d")
    except Exception:
        return datetime.now(timezone.utc).strftime("%Y-%m-%d")


def _extract_category(filename: str, paragraphs: list[str]) -> DocumentCategory:
    """Detects the document category from title block and body text."""
    body = (filename + " " + " ".join(paragraphs[:10])).lower()

    # Check explicit marker first
    for p in paragraphs[:5]:
        p_lower = p.lower()
        if "standard operating procedure" in p_lower or "sop" in p_lower:
            return DocumentCategory.SOP

    if any(k in body for k in ["policy", "guidelines", "regulation", "rules"]):
        return DocumentCategory.POLICY
    if any(k in body for k in ["circular", "notice", "notification", "memo"]):
        return DocumentCategory.CIRCULAR
    if any(k in body for k in ["handbook", "manual", "guidebook"]):
        return DocumentCategory.HANDBOOK
    if any(k in body for k in ["accreditation", "naac", "nba", "iqac"]):
        return DocumentCategory.ACCREDITATION
    if any(k in body for k in ["guideline", "guide", "framework"]):
        return DocumentCategory.GUIDELINE

    # VIT corpus default: all are SOPs
    return DocumentCategory.SOP


def _extract_access_level(paragraphs: list[str]) -> AccessLevel:
    """Suggests access level from first-page body markers. Defaults to Public."""
    body = " ".join(paragraphs[:15]).lower()
    if any(k in body for k in ["admin only", "confidential", "restricted"]):
        return AccessLevel.ADMIN
    if any(k in body for k in ["faculty only", "for faculty"]):
        return AccessLevel.FACULTY
    if any(k in body for k in ["student only", "for students only"]):
        return AccessLevel.STUDENT
    return AccessLevel.PUBLIC


# ===========================================================================
# SECTION 3: Table-Aware Text Extraction
# ===========================================================================

# Field labels that appear in the left column of VIT SOP tables
SOP_TABLE_LABELS = {
    "key objectives", "objectives", "key inputs", "inputs",
    "process description", "key outputs", "outputs",
    "key performers", "performers", "records", "documents",
    "pre-requisites", "prerequisites", "key performance", "sub process",
}

# Maps source table labels to the canonical serialization field names
LABEL_CANONICAL = {
    "key objectives": "Objectives",
    "objectives": "Objectives",
    "key inputs": "Inputs",
    "inputs": "Inputs",
    "process description": "Process Description",
    "key outputs": "Outputs",
    "outputs": "Outputs",
    "key performers": "Performers",
    "performers": "Performers",
    "records": "Records",
    "documents": "Documents",
    "pre-requisites": "Pre-requisites",
    "prerequisites": "Pre-requisites",
    "key performance": "Key Performance",
}

# Labels that identify a sub-process header cell
SUB_PROCESS_MARKERS = re.compile(
    r'^\d+[\.\d]*\s*:?\s*sub\s*process', re.IGNORECASE
)

def _is_sub_process_header(cell_text: str) -> bool:
    """Returns True if a cell looks like a sub-process title row."""
    return bool(SUB_PROCESS_MARKERS.match(cell_text.strip()))


def _is_label_cell(cell_text: str) -> bool:
    """Returns True if a cell is a known SOP field label (left column)."""
    return cell_text.strip().lower().rstrip(":") in SOP_TABLE_LABELS


def _extract_sub_process_name(header_text: str) -> str:
    """Extracts the human-readable sub-process name from a header cell."""
    text = header_text.strip()
    # e.g. "1.1 Sub Process: Conducting Term Test" → "Conducting Term Test"
    m = re.search(r'sub\s*process\s*:?\s*(.+)$', text, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    # e.g. "1.1 : Admissions Committee" → "Admissions Committee"
    m = re.search(r'^\d+[\.\d]*\s*:?\s*(.+)$', text)
    if m:
        return m.group(1).strip()
    return text


def _canonical_label(raw_label: str) -> str:
    """Maps a source table label to the canonical serialization field name."""
    key = raw_label.strip().lower().rstrip(":")
    return LABEL_CANONICAL.get(key, raw_label.strip().rstrip(":"))


def _serialize_sop_table(table) -> str:
    """
    Serializes a VIT SOP sub-process table into labeled text blocks.

    Produces output like:
        [SUB PROCESS]
        Name: Conducting Term Test

        Objectives:
        To appoint a committee...

        Inputs:
        DTE/University rules...

        Process Description:
        The committee shall be constituted...
    """
    lines: list[str] = []
    seen_cells: set[str] = set()
    current_label: str = ""

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        unique_cells = list(dict.fromkeys(c for c in cells if c))

        if not unique_cells:
            continue

        first = unique_cells[0]

        if _is_sub_process_header(first):
            if lines:
                lines.append("")
            lines.append("[SUB PROCESS]")
            name = _extract_sub_process_name(first)
            lines.append(f"Name: {name}")
            current_label = ""
            seen_cells = {first, name}

        elif _is_label_cell(first) and len(unique_cells) >= 2:
            label = _canonical_label(first)
            value = " | ".join(c for c in unique_cells[1:] if c not in seen_cells)
            if value:
                lines.append(f"{label}:")
                lines.append(value)
                seen_cells.update(unique_cells)
            current_label = label

        else:
            for c in unique_cells:
                if c not in seen_cells and len(c) > 3:
                    if current_label:
                        lines.append(c)
                    else:
                        lines.append(c)
                    seen_cells.add(c)

    return "\n".join(lines).strip()


def _serialize_generic_table(table) -> str:
    """
    Serializes a non-SOP table (e.g. pre-requisites, ToC) as plain text rows.
    Deduplicates merged cells within each row.
    """
    lines: list[str] = []
    seen_cells: set[str] = set()
    for row in table.rows:
        cells = list(dict.fromkeys(
            cell.text.strip() for cell in row.cells if cell.text.strip()
        ))
        row_text = " | ".join(c for c in cells if c and c not in seen_cells)
        if row_text:
            lines.append(row_text)
        seen_cells.update(cells)
    return "\n".join(lines)


def _has_sop_structure(table) -> bool:
    """Returns True if this table follows the VIT SOP sub-process layout."""
    for row in table.rows[:5]:
        for cell in row.cells:
            if _is_sub_process_header(cell.text) or _is_label_cell(cell.text):
                return True
    return False


def _extract_text_from_document(doc: docx.Document) -> tuple[list[str], str]:
    """
    Extracts and serializes all content from a python-docx Document object.

    Returns:
        (paragraph_texts, serialized_full_text)
        - paragraph_texts: list of raw non-empty paragraph strings (for metadata extraction)
        - serialized_full_text: full cleaned text suitable for chunking
    """
    # --- Paragraph extraction with heading tracking ---
    paragraph_texts: list[str] = []
    body_blocks: list[str] = []
    toc_pattern = re.compile(r'.+\t\d+\s*$')  # e.g. "1. Process: Admissions\t3"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Skip ToC lines
        if toc_pattern.match(text):
            continue

        paragraph_texts.append(text)

        # Mark heading-style paragraphs
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            body_blocks.append(f"[PROCESS: {text}]")
        else:
            body_blocks.append(text)

    # --- Table extraction ---
    table_blocks: list[str] = []
    for table in doc.tables:
        if _has_sop_structure(table):
            serialized = _serialize_sop_table(table)
        else:
            serialized = _serialize_generic_table(table)
        if serialized:
            table_blocks.append(serialized)

    # Combine: paragraphs first, then tables
    # Note: In VIT docs, tables contain the substantive content.
    # We place tables after paragraphs so heading context is established first.
    all_blocks = body_blocks + [""] + table_blocks
    full_text = "\n\n".join(b for b in all_blocks if b)
    return paragraph_texts, full_text


def _extract_text_from_pdf(filepath: str) -> tuple[list[str], str]:
    """
    Extracts text from a PDF using pypdf, mirroring the return contract of
    _extract_text_from_document so downstream metadata extraction and chunking
    are format-agnostic.

    Returns:
        (paragraph_texts, serialized_full_text)
    """
    from pypdf import PdfReader

    reader = PdfReader(filepath)
    paragraph_texts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        for line in page_text.split("\n"):
            line = line.strip()
            if line:
                paragraph_texts.append(line)

    full_text = "\n\n".join(paragraph_texts)
    return paragraph_texts, full_text


# ===========================================================================
# SECTION 4: Text Cleaning
# ===========================================================================

UNICODE_REPLACEMENTS = {
    "\u2013": "-",    # en dash
    "\u2014": "--",   # em dash
    "\u2019": "'",    # right single quotation mark
    "\u201c": '"',    # left double quotation mark
    "\u201d": '"',    # right double quotation mark
    "\u2022": "*",    # bullet
    "\u00a0": " ",    # non-breaking space
    "\ufffd": "",     # Unicode replacement character (corrupt byte)
    "\u2026": "...",  # ellipsis
}

def clean_text(text: str) -> str:
    """
    Cleans extracted text:
      1. Replaces known Unicode artifacts with ASCII equivalents.
      2. Collapses runs of 3+ blank lines to 2 blank lines.
      3. Strips trailing whitespace from each line.
    """
    for char, replacement in UNICODE_REPLACEMENTS.items():
        text = text.replace(char, replacement)

    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Strip trailing whitespace per line
    text = "\n".join(line.rstrip() for line in text.split("\n"))

    return text.strip()


# ===========================================================================
# SECTION 5: Incremental Ingestion & Version Supersession
# ===========================================================================

def _get_existing_hash(source_file: str) -> Optional[str]:
    """Returns the stored SHA-256 hash of a previously ingested document, or None."""
    record = ledger.get_document_by_source(source_file)
    return record["sha256_hash"] if record else None


def _check_version_supersession(department: str, new_version: str, new_doc_id: str) -> Optional[str]:
    """
    Detects if an older version of the same department module exists.

    Returns:
        The doc_id of the old document if supersession is needed, else None.
    """
    existing = ledger.get_document_by_department(department)
    if not existing:
        return None
    if existing["doc_id"] == new_doc_id:
        return None  # Same document — no supersession
    if existing.get("status") == DocumentStatus.SUPERSEDED.value:
        return None  # Already superseded

    # Compare versions: 'Final' and '3.0' always supersede '1.0'
    old_ver = existing.get("version", "1.0")
    if old_ver == new_version:
        return None  # Same version — treat as update, not supersession

    logger.info(
        f"[INGESTION] Version supersession: '{department}' "
        f"v{old_ver} → v{new_version}"
    )
    return existing["doc_id"]


# ===========================================================================
# SECTION 6: Main Ingestion Pipeline
# ===========================================================================

def load_staging_documents() -> list[str]:
    """Returns sorted list of absolute paths to all .docx files in data/staging/."""
    if not os.path.exists(STAGING_DIR):
        logger.error(f"[INGESTION] Staging directory not found: {STAGING_DIR}")
        return []
    return sorted(
        os.path.join(STAGING_DIR, f)
        for f in os.listdir(STAGING_DIR)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    )


def ingest_document(filepath: str, summary: IngestionSummary) -> Optional[tuple[DocumentRecord, str]]:
    """
    Ingests a single staged .docx file:
      1. Hashes the file for deduplication.
      2. Skips if hash is unchanged (incremental mode).
      3. Extracts and serializes text (paragraphs + tables).
      4. Builds DocumentRecord with full metadata.
      5. Registers to SQLite ledger.
      6. Returns (DocumentRecord, cleaned_text) for the chunker.

    Returns None if skipped or failed.
    """
    rel_path = os.path.relpath(filepath, PROJECT_ROOT)
    filename = os.path.basename(filepath)

    try:
        file_hash = compute_sha256(filepath)
    except Exception as e:
        logger.error(f"[INGESTION] Hash failed for '{filename}': {e}")
        summary.total_docs_failed += 1
        summary.errors.append(f"Hash failed: {filename}: {e}")
        return None

    # --- Incremental check: skip if hash unchanged ---
    existing_hash = _get_existing_hash(rel_path)
    if existing_hash and existing_hash == file_hash:
        logger.info(f"[INGESTION] SKIP (unchanged): {filename}")
        summary.total_docs_skipped += 1
        return None

    logger.info(f"[INGESTION] Processing: {filename}")

    ext = os.path.splitext(filename)[1].lower()

    # --- Open + Text Extraction (format-aware) ---
    try:
        if ext == ".pdf":
            paragraph_texts, raw_text = _extract_text_from_pdf(filepath)
        else:
            doc = docx.Document(filepath)
            paragraph_texts, raw_text = _extract_text_from_document(doc)
        cleaned_text = clean_text(raw_text)
    except Exception as e:
        logger.error(f"[INGESTION] Extraction failed for '{filename}': {e}")
        summary.total_docs_failed += 1
        summary.errors.append(f"Extraction failed: {filename}: {e}")
        ledger.log_event(rel_path, "ingest_extract", "failed", str(e))
        return None

    if len(cleaned_text) < 100:
        logger.warning(f"[INGESTION] Very short content for '{filename}' ({len(cleaned_text)} chars). Flagging.")

    # --- Metadata Extraction ---
    department  = _extract_module_name(filename, paragraph_texts)
    title       = _extract_title(paragraph_texts)
    category    = _extract_category(filename, paragraph_texts)
    version     = _extract_version(filename, paragraph_texts)
    upload_date = _extract_date(filepath, paragraph_texts)
    access_level = _extract_access_level(paragraph_texts)

    doc_record = DocumentRecord(
        doc_id       = file_hash,
        source_file  = rel_path,
        original_file= rel_path.replace("data\\staging\\", "data\\").replace("data/staging/", "data/"),
        title        = title,
        category     = category,
        department   = department,
        version      = version,
        access_level = access_level,
        upload_date  = upload_date,
        status       = DocumentStatus.INGESTED,
    )

    # --- Version supersession detection ---
    superseded_id = _check_version_supersession(department, version, file_hash)
    if superseded_id:
        ledger.mark_document_superseded(superseded_id)
        summary.total_docs_superseded += 1

    # --- Register in SQLite ledger ---
    try:
        ledger.upsert_document(doc_record.to_ledger_dict())
        ledger.log_event(rel_path, "ingest", "success", f"Ingested '{filename}' as {category.value}")
    except Exception as e:
        logger.error(f"[INGESTION] Ledger write failed for '{filename}': {e}")
        summary.total_docs_failed += 1
        summary.errors.append(f"Ledger write failed: {filename}: {e}")
        return None

    summary.total_docs_processed += 1
    return doc_record, cleaned_text


def run_ingestion(staging_dir: Optional[str] = None) -> IngestionSummary:
    """
    Main orchestrator for Phase 2 + Phase 3.

    1. Discovers all staged .docx files.
    2. Ingests each document (hash check → extract → metadata → ledger).
    3. Passes all (doc_record, text) pairs to the chunker.
    4. Returns IngestionSummary with full run statistics.

    This function is designed for both:
      - Full initial ingestion of all 20 VIT SOP documents.
      - Incremental ingestion of newly uploaded documents (admin workflow).
    """
    ledger.initialize_db()

    run_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    summary = IngestionSummary(
        run_id     = run_id,
        started_at = datetime.now(timezone.utc).isoformat(),
    )

    # --- Discover staged files ---
    staging_path = staging_dir or STAGING_DIR
    files = sorted(
        os.path.join(staging_path, f)
        for f in os.listdir(staging_path)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    )
    summary.total_docs_discovered = len(files)
    logger.info(f"[INGESTION] Run {run_id}: Found {len(files)} staged documents.")

    # --- Ingest each document ---
    doc_records: list[DocumentRecord] = []
    doc_texts:   dict[str, str]       = {}

    for filepath in files:
        result = ingest_document(filepath, summary)
        if result:
            doc_record, cleaned_text = result
            doc_records.append(doc_record)
            doc_texts[doc_record.doc_id] = cleaned_text

    # --- Phase 3: Chunk all ingested documents ---
    if doc_records:
        logger.info(f"[INGESTION] Passing {len(doc_records)} documents to chunker...")
        all_chunks = run_chunking(doc_records, doc_texts)
        summary.total_chunks_generated = sum(len(v) for v in all_chunks.values())
    else:
        all_chunks = {}

    summary.completed_at = datetime.now(timezone.utc).isoformat()

    # --- Print summary ---
    print(summary.report())
    ledger.log_event(
        "system",
        f"ingestion_run_{run_id}",
        "complete",
        f"Processed={summary.total_docs_processed}, "
        f"Skipped={summary.total_docs_skipped}, "
        f"Chunks={summary.total_chunks_generated}"
    )

    return summary


# ===========================================================================
# Entry Point
# ===========================================================================

if __name__ == "__main__":
    logger.info("Starting Phase 2 + 3: Document Ingestion & Chunking...")
    summary = run_ingestion()
    sys.exit(0 if summary.total_docs_failed == 0 else 1)
