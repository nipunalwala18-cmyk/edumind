import os
import sys
import hashlib
import json
import shutil
import re
from pathlib import Path
from datetime import datetime

# Import database ledger module
try:
    import ledger
except ImportError:
    # Fallback to local path import
    sys.path.append(os.path.dirname(os.path.abspath(__file__)))
    import ledger

try:
    import win32com.client
    import pythoncom
except ImportError:
    print("CRITICAL ERROR: pywin32 is not installed or not configured in your virtual environment.")
    print("Please install pywin32 using: pip install pywin32")
    sys.exit(1)

try:
    import docx
except ImportError:
    print("CRITICAL ERROR: python-docx is not installed in your virtual environment.")
    print("Please install python-docx using: pip install python-docx")
    sys.exit(1)

# Paths Configuration
PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
DATA_DIR = os.path.join(PROJECT_ROOT, "data")
STAGING_DIR = os.path.join(DATA_DIR, "staging")

# Initialize SQLite database
ledger.initialize_db()

def compute_sha256(filepath):
    """Computes SHA-256 hash of a file."""
    sha256 = hashlib.sha256()
    with open(filepath, "rb") as f:
        while chunk := f.read(8192):
            sha256.update(chunk)
    return sha256.hexdigest()

def normalize_filename(name):
    """Normalizes document filename for version group matching."""
    name = name.lower()
    # Remove prefix numbers e.g. "1. ", "10."
    name = re.sub(r'^\d+\s*[\.\-]\s*', '', name)
    # Remove "vit"
    name = re.sub(r'\bvit\b', '', name)
    # Remove version indicators
    name = re.sub(r'\b\d+\.\d+\b', '', name)
    name = re.sub(r'\bv\d+(\.\d+)?\b', '', name)
    name = re.sub(r'\(.*?\)', '', name)
    name = name.replace('final', '')
    # Remove file extensions
    name = re.sub(r'\.docx?(\.docx?)?$', '', name)
    # Cleanup spacing and special symbols
    name = re.sub(r'[\s_\-\.\(\)]+', ' ', name).strip()
    return name

def extract_date(filepath, text):
    """Extracts date from text, falling back to file modification time."""
    patterns = [
        r'\b\d{4}[-/]\d{2}[-/]\d{2}\b', # YYYY-MM-DD
        r'\b\d{2}[-/]\d{2}[-/]\d{4}\b', # DD-MM-YYYY or MM-DD-YYYY
        r'\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b', # DD Month YYYY
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b', # Month DD, YYYY
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            return matches[0]
            
    # Fallback to filesystem last modified date
    try:
        mtime = os.path.getmtime(filepath)
        return datetime.fromtimestamp(mtime).strftime('%Y-%m-%d')
    except Exception:
        return ""

def extract_version(filename, text):
    """Suggests document version based on filename and contents."""
    # Check filename first for pattern like 1.0 or v2.0 or (3) or (1)
    match = re.search(r'\b(?:v|version)?\s*(\d+\.\d+)\b', filename, re.IGNORECASE)
    if match:
        return match.group(1)
        
    match = re.search(r'\((\d+)\)', filename)
    if match:
        return f"{match.group(1)}.0"
    
    # Check text
    match = re.search(r'\b(?:version|ver\.|rev\.)\s*(\d+\.\d+)\b', text, re.IGNORECASE)
    if match:
        return match.group(1)
        
    if "final" in filename.lower():
        return "Final"
        
    return "1.0" # Default version

def extract_department(filename):
    """Suggests issuing department/module based on normalized filename."""
    # Clean filename to get raw module name
    clean = filename
    # Remove prefix numbers e.g. "1. ", "10."
    clean = re.sub(r'^\d+\s*[\.\-]\s*', '', clean)
    # Remove "vit"
    clean = re.sub(r'\bvit\b', '', clean, flags=re.IGNORECASE)
    # Remove version indicators
    clean = re.sub(r'\b\d+\.\d+\b', '', clean)
    clean = re.sub(r'\bv\d+(\.\d+)?\b', '', clean, flags=re.IGNORECASE)
    clean = re.sub(r'\(.*?\)', '', clean)
    clean = re.sub(r'\bfinal\b', '', clean, flags=re.IGNORECASE)
    # Remove extensions
    clean = re.sub(r'\.docx?(\.docx?)?$', '', clean, flags=re.IGNORECASE)
    # Clean up dm or HR suffixes
    clean = re.sub(r'dm$', '', clean, flags=re.IGNORECASE)
    clean = re.sub(r'_hr$', '', clean, flags=re.IGNORECASE)
    # Cleanup spacing and special symbols
    clean = re.sub(r'[\s_\-\.\(\)]+', ' ', clean).strip()
    return clean.title()

def suggest_category(filename, text):
    """Suggests document category based on filename and header text."""
    content = (filename + " " + text[:1000]).lower()
    if any(k in content for k in ["standard operating procedure", "sop", "procedure"]):
        return "SOP"
    elif any(k in content for k in ["policy", "guidelines", "rules", "regulation"]):
        return "Policy"
    elif any(k in content for k in ["circular", "notice", "memo", "notification"]):
        return "Circular"
    elif any(k in content for k in ["handbook", "manual", "guidebook"]):
        return "Handbook"
    elif any(k in content for k in ["accreditation", "nba", "naac", "iqac", "autonomous"]):
        return "Accreditation"
    return "SOP" # Fallback default for VIT documents

def extract_access_level(text):
    """Suggests access control tier based on text markers."""
    content = text[:2000].lower()
    if "admin only" in content or "confidential" in content or "restricted access" in content:
        return "Admin"
    elif "faculty only" in content or "faculty members" in content:
        return "Faculty"
    elif "student only" in content or "students only" in content:
        return "Student"
    return "Public" # Default fallback

def convert_doc_to_docx(abs_doc_path, abs_docx_path, word_app):
    """Converts a legacy .doc file to modern .docx format via Microsoft Word COM automation."""
    src = os.path.normpath(abs_doc_path)
    dest = os.path.normpath(abs_docx_path)
    
    doc = word_app.Documents.Open(src)
    doc.SaveAs2(dest, FileFormat=16) # FileFormat=16 is for wdFormatXMLDocument (.docx)
    doc.Close()

def run_assessment():
    print("=" * 60)
    print("STARTING PHASE 1: REPOSITORY ASSESSMENT & AUDITING")
    print(f"Project Directory: {PROJECT_ROOT}")
    print(f"Data Directory: {DATA_DIR}")
    print("=" * 60)

    # Ensure staging folder exists and is clean
    if os.path.exists(STAGING_DIR):
        shutil.rmtree(STAGING_DIR)
    os.makedirs(STAGING_DIR, exist_ok=True)

    # Lists for stats tracking
    total_files = 0
    file_types = {}
    duplicates_map = {}
    large_documents = []
    corrupted_files = []
    empty_documents = []
    doc_infos = []

    # 1. Walk directory and catalog raw files
    raw_files = []
    for root, dirs, files in os.walk(DATA_DIR):
        if "staging" in Path(root).parts:
            continue
        for file in files:
            if file.startswith("~$"):
                continue
            
            filepath = os.path.join(root, file)
            raw_files.append(filepath)

    total_files = len(raw_files)
    print(f"Found {total_files} raw candidate files in repository.")

    # 2. Hashing files to identify binary duplicates
    for filepath in raw_files:
        ext = os.path.splitext(filepath)[1].lower().replace(".", "")
        if not ext:
            ext = "no_extension"
        file_types[ext] = file_types.get(ext, 0) + 1
        
        file_size = os.path.getsize(filepath)
        if file_size > 300 * 1024:  # 300 KB threshold
            large_documents.append({
                "filename": os.path.basename(filepath),
                "size_bytes": file_size,
                "path": os.path.relpath(filepath, PROJECT_ROOT)
            })

        try:
            file_hash = compute_sha256(filepath)
            duplicates_map.setdefault(file_hash, []).append(filepath)
        except Exception as e:
            corrupted_files.append({
                "filename": os.path.basename(filepath),
                "error": f"Hashing failed: {str(e)}",
                "path": os.path.relpath(filepath, PROJECT_ROOT)
            })

    # 3. Headless COM automation loop for conversions
    print("Initializing headless MS Word automation instance...")
    word_app = None
    try:
        pythoncom.CoInitialize()
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0 # wdAlertsNone
    except Exception as e:
        print(f"CRITICAL ERROR: Failed to launch MS Word COM instance: {str(e)}")
        sys.exit(1)

    try:
        for filepath in raw_files:
            if filepath in [c["path"] for c in corrupted_files]:
                continue
                
            filename = os.path.basename(filepath)
            ext = os.path.splitext(filename)[1].lower()
            rel_path = os.path.relpath(filepath, DATA_DIR)
            
            # Clean double extension and setup staging path
            clean_filename = re.sub(r'\.docx?(\.docx?)?$', '.docx', filename)
            staging_path = os.path.join(STAGING_DIR, clean_filename)
            
            print(f"Processing: {filename} ...")
            try:
                if ext == ".doc":
                    convert_doc_to_docx(filepath, staging_path, word_app)
                    ledger.log_event(rel_path, "convert_doc", "success", f"Converted {filename} to staging docx.")
                elif ext in [".docx", ".docx.docx"]:
                    shutil.copy2(filepath, staging_path)
                    ledger.log_event(rel_path, "copy_docx", "success", f"Copied {filename} to staging docx.")
                else:
                    continue
            except Exception as e:
                print(f" -> ERROR: Conversion failed for {filename}: {str(e)}")
                corrupted_files.append({
                    "filename": filename,
                    "error": f"COM Conversion error: {str(e)}",
                    "path": os.path.relpath(filepath, PROJECT_ROOT)
                })
                ledger.log_event(rel_path, "convert_or_copy", "failed", str(e))
    finally:
        # Guarantee Word quits
        if word_app is not None:
            try:
                word_app.Quit()
                print("MS Word automation closed successfully.")
            except Exception:
                pass

    # 4. Text Inspection & Ingestion Ledger Write
    print("Reading staging files and analyzing content structure...")
    for file in os.listdir(STAGING_DIR):
        if not file.lower().endswith(".docx"):
            continue
            
        staging_filepath = os.path.join(STAGING_DIR, file)
        orig_rel_path = ""
        
        # Locate corresponding original filepath
        for rf in raw_files:
            rf_name = os.path.basename(rf)
            clean_rf_name = re.sub(r'\.docx?(\.docx?)?$', '.docx', rf_name)
            if clean_rf_name == file:
                orig_rel_path = os.path.relpath(rf, DATA_DIR)
                break
                
        if not orig_rel_path:
            orig_rel_path = file

        try:
            doc_read = docx.Document(staging_filepath)
            
            # Extract paragraphs text
            paragraphs_text = []
            for p in doc_read.paragraphs:
                txt = p.text.strip()
                if txt:
                    paragraphs_text.append(txt)
                    
            # Extract tables text
            tables_text = []
            for table in doc_read.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt and txt not in tables_text:
                            tables_text.append(txt)
                            
            full_text = "\n".join(paragraphs_text + tables_text)
            char_count = len(full_text)
            
            # Empty check
            if char_count < 100:
                empty_documents.append({
                    "filename": file,
                    "char_count": char_count,
                    "path": orig_rel_path
                })
            
            # Suggest Metadata
            title = paragraphs_text[0] if len(paragraphs_text) > 0 else file
            if len(title) > 100:
                title = title[:97] + "..."
                
            category = suggest_category(file, full_text)
            dept = extract_department(file)
            version = extract_version(file, full_text)
            doc_date = extract_date(staging_filepath, full_text)
            access = extract_access_level(full_text)
            
            doc_meta = {
                "title": title,
                "category": category,
                "department": dept,
                "version": version,
                "date": doc_date,
                "access_level": access
            }
            
            # Register in SQLite database
            file_hash = compute_sha256(staging_filepath)
            ledger.register_document(
                filepath=orig_rel_path,
                sha256_hash=file_hash,
                status="assessed",
                metadata=doc_meta
            )
            
            doc_infos.append({
                "filename": file,
                "original_path": orig_rel_path,
                "hash": file_hash,
                "character_count": char_count,
                "paragraphs": len(doc_read.paragraphs),
                "tables": len(doc_read.tables),
                "metadata": doc_meta
            })
            
        except Exception as e:
            print(f" -> ERROR: Content extraction failed for {file}: {str(e)}")
            corrupted_files.append({
                "filename": file,
                "error": f"Parsing/Read error: {str(e)}",
                "path": orig_rel_path
            })
            ledger.log_event(orig_rel_path, "parse_docx", "failed", str(e))

    # 5. Process Duplicates mapping
    duplicates_report = []
    for f_hash, paths in duplicates_map.items():
        if len(paths) > 1:
            duplicates_report.append([os.path.basename(p) for p in paths])

    # 6. Group potential version sequences
    version_groups = {}
    for info in doc_infos:
        base_name = normalize_filename(info["filename"])
        version_groups.setdefault(base_name, []).append({
            "filename": info["filename"],
            "version": info["metadata"]["version"],
            "date": info["metadata"]["date"]
        })
    # Filter version groups to only show sequences where we have version tokens
    # e.g., checking if version != '1.0' or group size > 1
    version_groups_filtered = {k: v for k, v in version_groups.items() if len(v) > 1 or any(x["version"] != "1.0" for x in v)}

    # 7. Categorization distributions
    categories_stats = {}
    for info in doc_infos:
        cat = info["metadata"]["category"]
        categories_stats[cat] = categories_stats.get(cat, 0) + 1

    departments_stats = {}
    for info in doc_infos:
        dep = info["metadata"]["department"]
        departments_stats[dep] = departments_stats.get(dep, 0) + 1

    # Format JSON report
    report_json = {
        "generated_at": datetime.now().isoformat(),
        "statistics": {
            "total_files_discovered": total_files,
            "total_files_assessed": len(doc_infos),
            "by_extension": file_types,
            "by_suggested_category": categories_stats,
            "by_suggested_department": departments_stats
        },
        "duplicate_groups": duplicates_report,
        "version_groups": version_groups_filtered,
        "large_documents": large_documents,
        "corrupted_documents": corrupted_files,
        "empty_documents": empty_documents,
        "document_inventory": doc_infos
    }

    # Write report.json
    with open(os.path.join(PROJECT_ROOT, "repository_report.json"), "w", encoding="utf-8") as f:
        json.dump(report_json, f, indent=2, ensure_ascii=False)
    print("Generated repository_report.json")

    # Generate Markdown report
    md_lines = [
        "# Repository Assessment & Auditing Report",
        f"\n**Report Generated On:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "\n---",
        "\n## 1. Repository Statistics",
        f"\n* **Total Files Discovered:** {total_files}",
        f"* **Total Files Successfully Assessed:** {len(doc_infos)}",
        f"* **Corrupted/Unreadable Files:** {len(corrupted_files)}",
        f"* **Empty/Low-Content Files:** {len(empty_documents)}",
        f"* **Duplicate Document Groups:** {len(duplicates_report)}",
        "\n### File Types Breakdown",
        "| Extension | Count |",
        "| --- | --- |"
    ]
    for ext, count in file_types.items():
        md_lines.append(f"| `.{ext}` | {count} |")

    md_lines.extend([
        "\n### Suggested Category Distribution",
        "| Category | Suggested Count |",
        "| --- | --- |"
    ])
    for cat, count in categories_stats.items():
        md_lines.append(f"| **{cat}** | {count} |")

    md_lines.extend([
        "\n### Represented Academic / Administrative Modules",
        "| Department / Module | Document Count |",
        "| --- | --- |"
    ])
    # Sort departments stats by count descending
    sorted_depts = sorted(departments_stats.items(), key=lambda x: x[1], reverse=True)
    for dep, count in sorted_depts:
        md_lines.append(f"| {dep} | {count} |")

    md_lines.extend([
        "\n---",
        "\n## 2. Integrity Issues Audited",
        "\n### Duplicate Files (Identical Checksums)"
    ])
    if duplicates_report:
        for idx, group in enumerate(duplicates_report):
            md_lines.append(f"{idx+1}. **Group {idx+1}**:")
            for item in group:
                md_lines.append(f"   * `{item}`")
    else:
        md_lines.append("*No binary duplicate files detected in this scan.*")

    md_lines.append("\n### Potential Version Suffixes & Series")
    if version_groups_filtered:
        for base, group in version_groups_filtered.items():
            md_lines.append(f"* **Module: {base.title()}**:")
            for file_item in sorted(group, key=lambda x: x["version"]):
                md_lines.append(f"  * File: `{file_item['filename']}` | Extracted Version: `{file_item['version']}` | Date: `{file_item['date']}`")
    else:
        md_lines.append("*No explicit version sequences detected via filename checks.*")

    md_lines.append("\n### Large Documents (> 300KB)")
    if large_documents:
        md_lines.append("| Filename | Size (KB) | Path |")
        md_lines.append("| --- | --- | --- |")
        for ld in large_documents:
            size_kb = round(ld['size_bytes'] / 1024, 1)
            md_lines.append(f"| `{ld['filename']}` | {size_kb} KB | `{ld['path']}` |")
    else:
        md_lines.append("*No files exceed the 300KB threshold.*")

    md_lines.append("\n### Empty / Low-Text Documents (< 100 characters)")
    if empty_documents:
        for ed in empty_documents:
            md_lines.append(f"* `{ed['filename']}` (Char count: {ed['char_count']})")
    else:
        md_lines.append("*No empty or low-content files detected.*")

    md_lines.append("\n### Corrupted or Conversion Failures")
    if corrupted_files:
        md_lines.append("| Filename | Error | Path |")
        md_lines.append("| --- | --- | --- |")
        for cf in corrupted_files:
            md_lines.append(f"| `{cf['filename']}` | {cf['error']} | `{cf['path']}` |")
    else:
        md_lines.append("*All files parsed and converted successfully without errors.*")

    md_lines.extend([
        "\n---",
        "\n## 3. Auto-Extracted Documents Inventory & Metadata",
        "\n| Staging Name | Category | Department / Module | Version | Source Date (FS Fallback) | Access Level |",
        "| --- | --- | --- | --- | --- | --- |"
    ])
    for info in doc_infos:
        meta = info["metadata"]
        md_lines.append(
            f"| `{info['filename']}` | **{meta['category']}** | {meta['department']} | "
            f"`{meta['version']}` | {meta['date'] or 'N/A'} | `{meta['access_level']}` |"
        )

    md_lines.extend([
        "\n---",
        "\n## 4. Suggested Metadata Extraction Rules",
        "\nBased on document text structure, we recommend the following rules for Phase 2 (Ingestion) extraction:",
        "\n1. **Category Mapping Rules**:",
        "   * If the first 3 lines contain the words `Standard Operating Procedure` or `SOP`, map as **SOP**.",
        "   * If containing `Policy` or `Guidelines` in the title block, map as **Policy**.",
        "   * Default fallback to folder categorization if matching keywords are absent.",
        "2. **Department Association Rules**:",
        "   * College SOPs represent specific administration modules. Map keyword filters based on first-page text headers (e.g. `Module: Admissions` maps to **Admissions** department).",
        "3. **Version & Revision Control Rules**:",
        "   * Search first page headers or tables for patterns matching `Version [0-9].[0-9]` or `Rev. [0-9].[0-9]` to extract standard version tags.",
        "   * Filenames with `1.0` or similar tags should be normalized, and only the latest verified version sequence should remain active in semantic search queries.",
        "4. **Security/Access Tiers**:",
        "   * Check first page text for keywords like `Confidential`, `Admin Only`, or `Faculty Only`.",
        "   * Propagate this level down to all chunks during the splitting phase."
    ])

    # Write report.md
    with open(os.path.join(PROJECT_ROOT, "repository_report.md"), "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    print("Generated repository_report.md")
    print("=" * 60)
    print("PHASE 1 COMPLETE. RUN REPORT AND DATABASE INITIALIZED SUCCESSFULLY.")
    print("=" * 60)

if __name__ == "__main__":
    run_assessment()
