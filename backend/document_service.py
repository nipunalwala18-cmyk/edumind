"""
backend/document_service.py
----------------------------
Secure document retrieval + rendering for the citation document viewer.

Responsibilities (kept separate from the FastAPI layer for testability):
  - RBAC: decide whether a role may open a document of a given access level.
    Public users are NEVER allowed to open documents.
  - Resolve a doc_id → physical file via the ledger (callers never see paths).
  - Path-traversal hardening: the resolved file must live under data/.
  - File-type detection (PDF / DOCX / DOC) by magic bytes, not just extension.
  - Render a self-contained, highlighted HTML viewer for the chunk that
    produced an answer.
  - Raw byte streaming (so endpoints stream files instead of exposing paths).

Public API:
    role_can_open(role)                         -> bool
    can_access(role, access_level)              -> bool
    load_document(doc_id)                        -> DocumentRef | None
    get_chunk_text(doc_id, chunk_index)          -> str
    iter_file_bytes(path)                         -> Iterator[bytes]
    render_viewer_html(docref, chunk_index)      -> str
    render_pending_preview_html(path, filename)  -> str
    media_type_for(path)                          -> str
"""

from __future__ import annotations

import base64
import html
import logging
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, Optional

logger = logging.getLogger(__name__)

_PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

# Only files under this root may ever be served (defence against traversal).
_DATA_ROOT = (Path(_PROJECT_ROOT) / "data").resolve()

# Max bytes we will base64-inline into a PDF viewer page (keeps the response
# self-contained without a second authenticated request). Larger PDFs fall back
# to the streaming /file endpoint.
_PDF_INLINE_CAP = 12 * 1024 * 1024  # 12 MB


# ---------------------------------------------------------------------------
# RBAC
# ---------------------------------------------------------------------------

# Which document access levels each role may open in the viewer.
# Public ROLE is intentionally absent → never allowed to open documents.
_ROLE_ACCESS: dict[str, set[str]] = {
    "Student": {"public", "student"},
    "Faculty": {"public", "student", "faculty"},
    "Admin":   {"public", "student", "faculty", "admin", "confidential"},
}

# Roles permitted to use the viewer at all.
_VIEWER_ROLES = frozenset({"Student", "Faculty", "Admin"})


def _norm_role(role: Optional[str]) -> str:
    return (role or "").strip().title()


def role_can_open(role: Optional[str]) -> bool:
    """True if the role is allowed to open documents at all (Public is not)."""
    return _norm_role(role) in _VIEWER_ROLES


def can_access(role: Optional[str], access_level: Optional[str]) -> bool:
    """True if `role` may open a document whose access level is `access_level`."""
    r = _norm_role(role)
    if r == "Admin":
        return True
    allowed = _ROLE_ACCESS.get(r)
    if not allowed:
        return False
    return (access_level or "Public").strip().lower() in allowed


# ---------------------------------------------------------------------------
# Document resolution (doc_id → file), traversal-hardened
# ---------------------------------------------------------------------------

@dataclass
class DocumentRef:
    doc_id:       str
    path:         Optional[Path] # Optional so we can hold reference even if file is missing
    title:        str
    department:   str
    version:      str
    category:     str
    access_level: str
    kind:         str          # "pdf" | "docx" | "doc" | "unknown"


def _safe_path(source_file: Optional[str]) -> Optional[Path]:
    """
    Resolves a ledger `source_file` to an absolute path and verifies it is
    inside the data/ root. Returns None on traversal attempts or missing files.
    """
    if not source_file:
        return None
    rel = str(source_file).replace("\\", "/")
    candidate = (Path(_PROJECT_ROOT) / rel).resolve()
    try:
        candidate.relative_to(_DATA_ROOT)
    except ValueError:
        logger.warning("[DOCVIEW] rejected out-of-root path: %s", source_file)
        return None
    return candidate if candidate.is_file() else None


def _safe_path_lenient(source_file: Optional[str]) -> Optional[Path]:
    """
    Resolves a ledger `source_file` to an absolute path and verifies it is
    inside the data/ root. Returns resolved path even if file is missing on disk.
    """
    if not source_file:
        return None
    rel = str(source_file).replace("\\", "/")
    candidate = (Path(_PROJECT_ROOT) / rel).resolve()
    try:
        candidate.relative_to(_DATA_ROOT)
    except ValueError:
        logger.warning("[DOCVIEW] rejected out-of-root path: %s", source_file)
        return None
    return candidate


def _detect_kind(path: Optional[Path], filename_fallback: str = "") -> str:
    """Detect file type by magic bytes, falling back to the extension."""
    if not path or not path.is_file():
        ext = os.path.splitext(filename_fallback or (path.name if path else ""))[1].lower().replace(".", "")
        return ext if ext in ("pdf", "docx", "doc") else "unknown"
    try:
        with open(path, "rb") as f:
            head = f.read(8)
    except OSError:
        head = b""
    if head[:4] == b"%PDF":
        return "pdf"
    if head[:4] == b"PK\x03\x04":          # zip container → docx
        return "docx"
    if head[:4] == b"\xD0\xCF\x11\xE0":    # OLE2 → legacy .doc
        return "doc"
    ext = path.suffix.lower().lstrip(".")
    return ext if ext in ("pdf", "docx", "doc") else "unknown"


def load_document(doc_id: str) -> Optional[DocumentRef]:
    """
    Looks up a document by doc_id in the ledger and returns a DocumentRef.
    Returns None if unknown, but does NOT return None if the file is missing from disk
    (sets path and kind appropriately so metadata can still be viewed).
    """
    if not doc_id:
        return None
    try:
        import ledger
        ledger.initialize_db()
        rec = ledger.get_document_by_doc_id(doc_id)
    except Exception as exc:
        logger.error("[DOCVIEW] ledger lookup failed for %s: %s", doc_id, exc)
        return None
    if not rec:
        return None

    source_file = rec.get("source_file") or rec.get("filepath") or ""
    path = _safe_path_lenient(source_file)
    if path is None:
        logger.warning("[DOCVIEW] path traversal or invalid path for doc_id=%s", doc_id)
        return None

    if not path.is_file():
        parent = path.parent
        if parent.is_dir():
            def _norm_name(n: str) -> str:
                return n.lower().replace(" ", "").replace(".docx", "").replace(".doc", "")
            target_norm = _norm_name(path.name)
            for child in parent.iterdir():
                if child.is_file() and _norm_name(child.name) == target_norm:
                    path = child
                    break

    return DocumentRef(
        doc_id=doc_id,
        path=path,
        title=rec.get("title") or path.stem,
        department=rec.get("department") or "General",
        version=rec.get("version") or "1.0",
        category=rec.get("category") or "SOP",
        access_level=rec.get("access_level") or "Public",
        kind=_detect_kind(path, source_file),
    )



# ---------------------------------------------------------------------------
# Chunk lookup (for highlighting)
# ---------------------------------------------------------------------------

def get_chunk_text(doc_id: str, chunk_index: Optional[int]) -> str:
    """Returns the stored text of a specific chunk (for highlighting)."""
    try:
        import ledger
        chunks = ledger.get_chunks_by_doc_id(doc_id)
    except Exception as exc:
        logger.error("[DOCVIEW] chunk lookup failed for %s: %s", doc_id, exc)
        return ""
    if not chunks:
        return ""
    if chunk_index is not None:
        for c in chunks:
            if c.get("chunk_index") == chunk_index:
                return c.get("content", "") or ""
    return chunks[0].get("content", "") or ""


# ---------------------------------------------------------------------------
# Streaming
# ---------------------------------------------------------------------------

_MEDIA = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc":  "application/msword",
}


def media_type_for(path: Path) -> str:
    return _MEDIA.get(path.suffix.lower(), "application/octet-stream")


def iter_file_bytes(path: Path, block: int = 64 * 1024) -> Iterator[bytes]:
    """Yields the file in blocks so endpoints stream rather than buffer."""
    with open(path, "rb") as f:
        while True:
            data = f.read(block)
            if not data:
                break
            yield data


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip().lower()


def _extract_docx_paragraphs(path: Path) -> list[str]:
    """Extracts non-empty paragraph + table text from a .docx file."""
    import docx
    doc = docx.Document(str(path))
    out: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            out.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                out.append(" | ".join(cells))
    return out


def render_docx_to_html(path: Path, chunk_text: str) -> tuple[str, bool]:
    """Converts a .docx file to a rich HTML layout (headings, tables, lists, styles) with highlighting."""
    import docx
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = docx.Document(str(path))
    chunk_norm = _norm(chunk_text)

    html_parts = []
    found_any = False

    for element in doc.element.body:
        if element.tag.endswith('p'):
            p = Paragraph(element, doc)
            text = p.text
            if not text or not text.strip():
                continue

            pnorm = _norm(text)
            is_hit = bool(chunk_norm) and len(pnorm) > 12 and pnorm in chunk_norm

            p_html = ""
            for run in p.runs:
                r_text = html.escape(run.text)
                if run.bold:
                    r_text = f"<strong>{r_text}</strong>"
                if run.italic:
                    r_text = f"<em>{r_text}</em>"
                p_html += r_text

            if not p_html and text:
                p_html = html.escape(text)

            anchor = ""
            hit_class = ""
            if is_hit:
                if not found_any:
                    anchor = ' id="hl"'
                hit_class = " hit"
                p_html = f"<mark>{p_html}</mark>"
                found_any = True

            style_name = p.style.name.lower()
            if "heading 1" in style_name:
                html_parts.append(f'<h1 class="heading-1{hit_class}"{anchor}>{p_html}</h1>')
            elif "heading 2" in style_name:
                html_parts.append(f'<h2 class="heading-2{hit_class}"{anchor}>{p_html}</h2>')
            elif "heading 3" in style_name:
                html_parts.append(f'<h3 class="heading-3{hit_class}"{anchor}>{p_html}</h3>')
            elif "list bullet" in style_name or "bullet" in style_name:
                html_parts.append(f'<li class="bullet-item{hit_class}"{anchor}>{p_html}</li>')
            else:
                html_parts.append(f'<p class="para{hit_class}"{anchor}>{p_html}</p>')

        elif element.tag.endswith('tbl'):
            t = Table(element, doc)
            table_html = ['<table class="doc-table">']
            for row in t.rows:
                table_html.append('<tr>')
                for cell in row.cells:
                    cell_html = []
                    for cp in cell.paragraphs:
                        if not cp.text or not cp.text.strip():
                            continue
                        cp_norm = _norm(cp.text)
                        is_cp_hit = bool(chunk_norm) and len(cp_norm) > 12 and cp_norm in chunk_norm

                        cp_text = ""
                        for run in cp.runs:
                            r_text = html.escape(run.text)
                            if run.bold:
                                r_text = f"<strong>{r_text}</strong>"
                            if run.italic:
                                r_text = f"<em>{r_text}</em>"
                            cp_text += r_text
                        if not cp_text and cp.text:
                            cp_text = html.escape(cp.text)

                        anchor = ""
                        if is_cp_hit:
                            if not found_any:
                                anchor = ' id="hl"'
                            cp_text = f"<mark>{cp_text}</mark>"
                            found_any = True
                        cell_html.append(f'<p class="table-para"{anchor}>{cp_text}</p>')
                    table_html.append(f'<td>{" ".join(cell_html)}</td>')
                table_html.append('</tr>')
            table_html.append('</table>')
            html_parts.append("\n".join(table_html))

    return "\n".join(html_parts), found_any


def _extract_pdf_paragraphs(path: Path) -> Optional[list[str]]:
    """Extracts text from a PDF if a PDF library is available, else None."""
    for module, fn in (
        ("pypdf", "_pdf_pypdf"),
        ("PyPDF2", "_pdf_pypdf"),
        ("pdfminer.high_level", "_pdf_pdfminer"),
    ):
        try:
            __import__(module)
        except Exception:
            continue
        try:
            return globals()[fn](path, module)
        except Exception as exc:           # pragma: no cover - lib-specific
            logger.warning("[DOCVIEW] PDF extract via %s failed: %s", module, exc)
    return None


def _pdf_pypdf(path: Path, module: str) -> list[str]:  # pragma: no cover
    mod = __import__(module)
    reader = mod.PdfReader(str(path))
    paras: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        paras.extend(p.strip() for p in text.split("\n") if p.strip())
    return paras


def _pdf_pdfminer(path: Path, module: str) -> list[str]:  # pragma: no cover
    from pdfminer.high_level import extract_text
    text = extract_text(str(path)) or ""
    return [p.strip() for p in text.split("\n") if p.strip()]


def _highlight_body(paragraphs: list[str], chunk_text: str) -> tuple[str, bool]:
    """
    Wraps every paragraph that belongs to the source chunk in <mark>. Returns
    (html_body, found_any). The first highlighted paragraph gets id="hl" so the
    viewer can scroll to it.
    """
    chunk_norm = _norm(chunk_text)
    parts: list[str] = []
    found = False
    for para in paragraphs:
        pnorm = _norm(para)
        is_hit = bool(chunk_norm) and len(pnorm) > 12 and pnorm in chunk_norm
        esc = html.escape(para)
        if is_hit:
            anchor = ' id="hl"' if not found else ""
            parts.append(f'<p class="para hit"{anchor}><mark>{esc}</mark></p>')
            found = True
        else:
            parts.append(f'<p class="para">{esc}</p>')
    return "\n".join(parts), found


_VIEWER_CSS = """
:root{--bg:#0a0a0b;--surface:#141417;--surface2:#1a1a1e;--border:rgba(255,255,255,.10);
--text:#f4f4f5;--text2:#a1a1aa;--text3:#6b6b75;--accent:#ffffff;--mark:#facc15;}
*{box-sizing:border-box}
body{margin:0;background:var(--bg);color:var(--text);
font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Inter,sans-serif;line-height:1.65;}
.doc-head{position:sticky;top:0;background:rgba(10,10,11,.92);backdrop-filter:blur(8px);
border-bottom:1px solid var(--border);padding:1rem 1.4rem;z-index:5;}
.doc-title{font-size:1.05rem;font-weight:700;}
.doc-meta{font-size:.78rem;color:var(--text3);margin-top:3px;}
.doc-flag{display:inline-block;margin-top:.5rem;font-size:.74rem;color:#0a0a0b;
background:var(--mark);border-radius:999px;padding:2px 10px;font-weight:600;}
.doc-body{max-width:820px;margin:0 auto;padding:1.6rem 1.4rem 4rem;}
.para{margin:0 0 .85rem;font-size:.93rem;color:var(--text);}
.para.hit{scroll-margin-top:90px;}
mark{background:var(--mark);color:#0a0a0b;padding:.05em .15em;border-radius:3px;}
.notice{background:var(--surface);border:1px solid var(--border);border-radius:12px;
padding:1rem 1.2rem;color:var(--text2);font-size:.9rem;margin-bottom:1.2rem;}
.chunk-box{background:var(--surface);border:1px solid var(--border);border-left:3px solid var(--mark);
border-radius:10px;padding:1rem 1.2rem;font-size:.9rem;white-space:pre-wrap;}
embed{width:100%;height:78vh;border:1px solid var(--border);border-radius:10px;background:#fff;}
.heading-1{font-size:1.6rem;font-weight:800;margin:1.8rem 0 1rem;color:var(--accent);border-bottom:1px solid var(--border);padding-bottom:0.4rem;}
.heading-1.hit{scroll-margin-top:90px;}
.heading-2{font-size:1.3rem;font-weight:700;margin:1.5rem 0 0.8rem;color:var(--text);}
.heading-2.hit{scroll-margin-top:90px;}
.heading-3{font-size:1.1rem;font-weight:600;margin:1.2rem 0 0.6rem;color:var(--text2);}
.heading-3.hit{scroll-margin-top:90px;}
.bullet-item{margin-left:1.5rem;margin-bottom:0.5rem;list-style-type:disc;font-size:0.93rem;color:var(--text);}
.bullet-item.hit{scroll-margin-top:90px;}
.doc-table{width:100%;border-collapse:collapse;margin:1.5rem 0;background:var(--surface);border:1px solid var(--border);border-radius:8px;overflow:hidden;}
.doc-table td{border:1px solid var(--border);padding:0.75rem 1rem;font-size:0.88rem;color:var(--text2);vertical-align:top;}
.doc-table tr:nth-child(even){background:rgba(255,255,255,0.02);}
.table-para{margin:0;font-size:0.88rem;line-height:1.4;}
"""

_SCROLL_JS = """
<script>
window.addEventListener('load',function(){
  var el=document.getElementById('hl');
  if(el){el.scrollIntoView({behavior:'smooth',block:'center'});}
});
</script>
"""


def _page(docref: DocumentRef, chunk_index: Optional[int], inner: str,
          flag: str = "Highlighted: the passage that generated the answer") -> str:
    cidx = "" if chunk_index is None else f" · source chunk {chunk_index}"
    return (
        "<!DOCTYPE html><html lang='en'><head><meta charset='UTF-8'>"
        "<meta name='viewport' content='width=device-width, initial-scale=1.0'>"
        f"<title>{html.escape(docref.title)}</title><style>{_VIEWER_CSS}</style></head><body>"
        "<div class='doc-head'>"
        f"<div class='doc-title'>{html.escape(docref.title)}</div>"
        f"<div class='doc-meta'>{html.escape(docref.department)} · v{html.escape(docref.version)}"
        f" · {html.escape(docref.category)}{cidx}</div>"
        f"<div class='doc-flag'>{html.escape(flag)}</div>"
        "</div>"
        f"<div class='doc-body'>{inner}</div>"
        f"{_SCROLL_JS}</body></html>"
    )


def render_viewer_html(docref: DocumentRef, chunk_index: Optional[int]) -> str:
    """
    Builds a self-contained, RBAC-cleared HTML viewer page with the source chunk
    highlighted. The page embeds everything it needs (no second authenticated
    request), so it can be shown directly in a sandboxed iframe.
    """
    chunk_text = get_chunk_text(docref.doc_id, chunk_index)

    # --- Fallback: File is missing from disk ---
    if docref.path is None or not docref.path.is_file():
        try:
            import ledger
            all_chunks = ledger.get_chunks_by_doc_id(docref.doc_id)
        except Exception:
            all_chunks = []
        
        adjacent_html = ""
        current_chunk_idx = chunk_index if chunk_index is not None else 0
        
        if all_chunks:
            adjacent_html += "<div style='margin-top: 1.5rem;'><h4 style='color: var(--text2); border-bottom: 1px solid var(--border); padding-bottom: 0.5rem;'>Document Chunk Navigation (Original File Unavailable)</h4>"
            adjacent_html += "<div style='display: flex; gap: 0.35rem; flex-wrap: wrap; margin-bottom: 1rem;'>"
            for c in all_chunks:
                c_idx = c.get("chunk_index", 0)
                active_style = "background: var(--mark); color: #0a0a0b; font-weight: 600;" if c_idx == current_chunk_idx else "background: var(--surface2); color: var(--text2);"
                adjacent_html += f"<a href='?chunk_index={c_idx}' style='text-decoration: none; padding: 4px 10px; border-radius: 4px; font-size: 0.8rem; {active_style}'>Chunk {c_idx + 1}</a>"
            adjacent_html += "</div>"
            
            active_chunk = None
            for c in all_chunks:
                if c.get("chunk_index") == current_chunk_idx:
                    active_chunk = c
                    break
            if not active_chunk and all_chunks:
                active_chunk = all_chunks[0]
                
            if active_chunk:
                content = active_chunk.get("content", "")
                section = active_chunk.get("section_heading", "")
                sect_info = f"<div style='font-size: 0.85rem; color: var(--text3); margin-bottom: 0.5rem;'>Section: {html.escape(section)}</div>" if section else ""
                adjacent_html += f"<div class='chunk-box' style='font-family: inherit; font-size: 0.93rem;'>{sect_info}{html.escape(content)}</div>"
            adjacent_html += "</div>"
        else:
            adjacent_html += _chunk_fallback(chunk_text, "Original source document is not available.")
            
        return _page(
            docref,
            chunk_index,
            f"<div class='notice' style='border-left: 3px solid #ef4444;'>Original document file is unavailable on disk. Displaying retrieved chunk evidence.</div>{adjacent_html}",
            flag="Chunk Fallback View (File Unavailable)"
        )

    # --- DOCX: rich document layout rendering + inline highlight ---
    if docref.kind == "docx":
        try:
            body, found = render_docx_to_html(docref.path, chunk_text)
            flag = ("Highlighted: the passage that generated the answer"
                    if found else "Document opened (exact passage not located)")
            return _page(docref, chunk_index, body, flag)
        except Exception as exc:
            logger.error("[DOCVIEW] docx render failed: %s", exc, exc_info=True)
            return _page(docref, chunk_index,
                         _chunk_fallback(chunk_text, "Could not render this document."))

    # --- PDF: text-highlight if a PDF lib exists, else inline-embed the file ---
    if docref.kind == "pdf":
        paras = _extract_pdf_paragraphs(docref.path)
        if paras:
            body, found = _highlight_body(paras, chunk_text)
            flag = ("Highlighted: the passage that generated the answer"
                    if found else "Document opened (exact passage not located)")
            return _page(docref, chunk_index, body, flag)
        # No PDF text library — embed the original PDF inline (if small enough).
        size = docref.path.stat().st_size
        if size <= _PDF_INLINE_CAP:
            b64 = base64.b64encode(docref.path.read_bytes()).decode("ascii")
            inner = (
                _chunk_fallback(chunk_text,
                                "Source passage (highlighted text shown above the document):")
                + f"<embed type='application/pdf' src='data:application/pdf;base64,{b64}'>"
            )
            return _page(docref, chunk_index, inner, "Source passage shown above · full PDF below")
        return _page(docref, chunk_index,
                     _chunk_fallback(chunk_text, "This PDF is large — passage shown below."))

    # --- DOC (legacy) / unknown: no safe in-browser renderer here ---
    return _page(
        docref, chunk_index,
        _chunk_fallback(
            chunk_text,
            "This is a legacy .doc file. The exact source passage is shown below; "
            "use “Download original” to open the full document.",
        ),
        flag="Source passage",
    )


def _chunk_fallback(chunk_text: str, message: str) -> str:
    safe = html.escape(chunk_text or "(passage unavailable)")
    return f"<div class='notice'>{html.escape(message)}</div><div class='chunk-box'>{safe}</div>"


def render_pending_preview_html(path: Path, filename: str) -> str:
    """
    Renders a not-yet-ingested (pending committee-head) submission for admin
    review, reusing the same docx/pdf renderers as the citation viewer but
    without a DocumentRef/chunk (nothing has been chunked or approved yet).
    """
    kind = _detect_kind(path, filename)
    title = html.escape(filename)

    if kind == "docx":
        try:
            body, _ = render_docx_to_html(path, "")
        except Exception as exc:
            logger.error("[DOCVIEW] pending docx render failed: %s", exc, exc_info=True)
            body = "<div class='notice'>Could not render this document.</div>"
    elif kind == "pdf":
        paras = _extract_pdf_paragraphs(path)
        if paras:
            body = "".join(f"<p class='para'>{html.escape(p)}</p>" for p in paras)
        else:
            size = path.stat().st_size
            if size <= _PDF_INLINE_CAP:
                b64 = base64.b64encode(path.read_bytes()).decode("ascii")
                body = f"<embed type='application/pdf' src='data:application/pdf;base64,{b64}'>"
            else:
                body = "<div class='notice'>This PDF is too large to preview inline. Download it to review.</div>"
    else:
        body = "<div class='notice'>This file type has no in-browser preview. Download it to review.</div>"

    return (
        "<!DOCTYPE html><html lang='en'><head><meta charset='UTF-8'>"
        "<meta name='viewport' content='width=device-width, initial-scale=1.0'>"
        f"<title>{title}</title><style>{_VIEWER_CSS}</style></head><body>"
        "<div class='doc-head'>"
        f"<div class='doc-title'>{title}</div>"
        "<div class='doc-flag'>Pending review — not yet in the knowledge base</div>"
        "</div>"
        f"<div class='doc-body'>{body}</div></body></html>"
    )


def get_chunk_by_id(chunk_id: str) -> Optional[dict]:
    """Retrieves a chunk's content and metadata by its chunk_id from the ledger."""
    try:
        import ledger
        conn = ledger.get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM chunks WHERE chunk_id = ?", (chunk_id,))
            row = cursor.fetchone()
            return dict(row) if row else None
        finally:
            conn.close()
    except Exception as exc:
        logger.error("[DOCVIEW] chunk lookup failed for chunk_id=%s: %s", chunk_id, exc)
        return None

